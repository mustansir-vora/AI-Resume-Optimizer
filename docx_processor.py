import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from lxml import etree
from io import BytesIO

# Namespace map for parsing DOCX XML
ns = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
}

def _get_run_properties_from_element(r_element):
    """Extracts all possible run properties from a w:r OxmlElement."""
    props = {
        "bold": "false", "italic": "false", "underline": "false",
        "strikethrough": "false", "font_name": "Unknown",
        "font_size": "0", "font_color": "#000000", "highlight_color": "none"
    }
    rPr = r_element.find('w:rPr', namespaces=ns)
    if rPr is None: return props

    if rPr.find('w:b', namespaces=ns) is not None: props['bold'] = "true"
    if rPr.find('w:i', namespaces=ns) is not None: props['italic'] = "true"
    if rPr.find('w:u', namespaces=ns) is not None: props['underline'] = "true"
    if rPr.find('w:strike', namespaces=ns) is not None: props['strikethrough'] = "true"
    
    rFonts = rPr.find('w:rFonts', namespaces=ns)
    if rFonts is not None and rFonts.get(qn('w:ascii')):
        props['font_name'] = rFonts.get(qn('w:ascii'))
    
    sz = rPr.find('w:sz', namespaces=ns)
    if sz is not None and sz.get(qn('w:val')):
        try:
            props['font_size'] = str(int(sz.get(qn('w:val'))) / 2)
        except (ValueError, TypeError): pass
        
    color = rPr.find('w:color', namespaces=ns)
    if color is not None and color.get(qn('w:val')):
        props['font_color'] = f"#{color.get(qn('w:val'))}"
        
    highlight = rPr.find('w:highlight', namespaces=ns)
    if highlight is not None and highlight.get(qn('w:val')):
        props['highlight_color'] = highlight.get(qn('w:val'))
        
    return props

def _get_paragraph_properties_from_element(p_element):
    """Extracts all possible paragraph properties from a w:p OxmlElement."""
    props = {
        "alignment": "left", "style": "Normal",
        "line_spacing": "0", "line_spacing_rule": "auto",
        "space_before": "0", "space_after": "0",
        "left_indent": "0", "right_indent": "0", "first_line_indent": "0",
        "list_type": "none", "list_level": "0",
        "shading_color": "auto", "shading_fill": "auto",
        "bottom_border_style": "none", "bottom_border_size": "0", "bottom_border_color": "auto"
    }
    pPr = p_element.find('w:pPr', namespaces=ns)
    if pPr is None: return props

    jc = pPr.find('w:jc', namespaces=ns)
    if jc is not None and jc.get(qn('w:val')):
        props['alignment'] = jc.get(qn('w:val'))

    pStyle = pPr.find('w:pStyle', namespaces=ns)
    if pStyle is not None and pStyle.get(qn('w:val')):
        props['style'] = pStyle.get(qn('w:val'))

    spacing = pPr.find('w:spacing', namespaces=ns)
    if spacing is not None:
        if spacing.get(qn('w:line')): props['line_spacing'] = spacing.get(qn('w:line'))
        if spacing.get(qn('w:lineRule')): props['line_spacing_rule'] = spacing.get(qn('w:lineRule'))
        if spacing.get(qn('w:before')): props['space_before'] = str(int(spacing.get(qn('w:before'))) / 20)
        if spacing.get(qn('w:after')): props['space_after'] = str(int(spacing.get(qn('w:after'))) / 20)

    ind = pPr.find('w:ind', namespaces=ns)
    if ind is not None:
        if ind.get(qn('w:left')): props['left_indent'] = str(int(ind.get(qn('w:left'))) / 20)
        if ind.get(qn('w:right')): props['right_indent'] = str(int(ind.get(qn('w:right'))) / 20)
        if ind.get(qn('w:firstLine')): props['first_line_indent'] = str(int(ind.get(qn('w:firstLine'))) / 20)

    numPr = pPr.find('w:numPr', namespaces=ns)
    if numPr is not None:
        props['list_type'] = 'numbered' # A simplification
        ilvl = numPr.find('w:ilvl', namespaces=ns)
        if ilvl is not None: props['list_level'] = ilvl.get(qn('w:val'))

    shd = pPr.find('w:shd', namespaces=ns)
    if shd is not None:
        if shd.get(qn('w:color')): props['shading_color'] = shd.get(qn('w:color'))
        if shd.get(qn('w:fill')): props['shading_fill'] = shd.get(qn('w:fill'))

    pBdr = pPr.find('w:pBdr', namespaces=ns)
    if pBdr is not None:
        bottom = pBdr.find('w:bottom', namespaces=ns)
        if bottom is not None:
            props['bottom_border_style'] = bottom.get(qn('w:val'))
            if bottom.get(qn('w:sz')): props['bottom_border_size'] = bottom.get(qn('w:sz'))
            if bottom.get(qn('w:color')): props['bottom_border_color'] = bottom.get(qn('w:color'))
            
    return props

def _process_paragraph_content(p_element, parent_xml, doc, images_data):
    """Processes runs, hyperlinks, and images within a paragraph."""
    for child in p_element:
        if child.tag == qn('w:r'):
            drawing = child.find('w:drawing', namespaces=ns)
            if drawing is not None:
                blip = drawing.find('.//a:blip', namespaces=ns)
                if blip is not None:
                    r_id = blip.get(qn('r:embed'))
                    if r_id and r_id in doc.part.rels:
                        image_part = doc.part.rels[r_id].target_part
                        images_data[r_id] = {
                            'bytes': image_part.blob,
                            'content_type': image_part.content_type
                        }
                        img_xml = etree.SubElement(parent_xml, "image")
                        img_xml.set("r_id", r_id)
                        drawing_str = etree.tostring(drawing).decode()
                        img_xml.set("drawing_xml", drawing_str)
                    continue

            r_xml = etree.SubElement(parent_xml, "run")
            run_props = _get_run_properties_from_element(child)
            for key, value in run_props.items(): r_xml.set(key, value)
            text = ''.join(t.text for t in child.findall('.//w:t', namespaces=ns) if t.text)
            text_element = etree.SubElement(r_xml, "text")
            text_element.text = etree.CDATA(text)

        elif child.tag == qn('w:hyperlink'):
            r_id = child.get(qn('r:id'))
            if not r_id or r_id not in doc.part.rels: continue
            url = doc.part.rels[r_id].target_ref
            hlink_xml = etree.SubElement(parent_xml, "hyperlink")
            hlink_xml.set("url", url)
            for run_element in child.findall('.//w:r', namespaces=ns):
                _process_paragraph_content(run_element, hlink_xml, doc, images_data)

def _process_body_elements(body_element, parent_xml, doc, images_data):
    """Recursively processes paragraphs and tables in the document body or a table cell."""
    for child in body_element:
        if child.tag == qn('w:p'):
            p_xml = etree.SubElement(parent_xml, "paragraph")
            p_props = _get_paragraph_properties_from_element(child)
            for key, value in p_props.items(): p_xml.set(key, str(value))
            _process_paragraph_content(child, p_xml, doc, images_data)
        elif child.tag == qn('w:tbl'):
            tbl_xml = etree.SubElement(parent_xml, "table_grid")
            for tr in child.findall('w:tr', namespaces=ns):
                row_xml = etree.SubElement(tbl_xml, "row")
                for tc in tr.findall('w:tc', namespaces=ns):
                    cell_xml = etree.SubElement(row_xml, "cell")
                    _process_body_elements(tc, cell_xml, doc, images_data)

def extract_docx_to_xml(docx_file_path):
    """
    Converts a DOCX file to a structured XML string, preserving content, styling, tables, and images.
    """
    if not os.path.exists(docx_file_path):
        raise FileNotFoundError(f"File not found: {docx_file_path}")

    doc = Document(docx_file_path)
    resume_xml = etree.Element("resume")
    images_data = {}

    _process_body_elements(doc.element.body, resume_xml, doc, images_data)

    xml_string = etree.tostring(resume_xml, pretty_print=True, xml_declaration=True, encoding='UTF-8').decode('utf-8')
    return xml_string, images_data

def _apply_run_properties(run, run_attrs):
    """Applies all captured run properties."""
    run.bold = run_attrs.get("bold") == "true"
    run.italic = run_attrs.get("italic") == "true"
    run.underline = run_attrs.get("underline") == "true"
    run.font.strike = run_attrs.get("strikethrough") == "true"
    
    if "font_name" in run_attrs and run_attrs['font_name'] != "Unknown":
        run.font.name = run_attrs["font_name"]
        
    if "font_size" in run_attrs and float(run_attrs.get("font_size", 0)) > 0:
        run.font.size = Pt(float(run_attrs["font_size"]))
        
    if "font_color" in run_attrs and run_attrs['font_color'] != "#000000":
        try:
            color = run_attrs["font_color"].lstrip('#')
            if len(color) == 6: run.font.color.rgb = RGBColor.from_string(color)
        except ValueError: pass
        
    if "highlight_color" in run_attrs and run_attrs['highlight_color'] != "none":
        run.font.highlight_color = getattr(WD_COLOR_INDEX, run_attrs['highlight_color'].upper(), None)

def _apply_paragraph_properties(p, para_attrs):
    """Applies all captured paragraph properties."""
    alignment_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
    p.alignment = alignment_map.get(para_attrs.get("alignment", "left").lower(), WD_ALIGN_PARAGRAPH.LEFT)

    try:
        p.style = para_attrs.get("style", "Normal")
    except Exception:
        p.style = "Normal"

    pf = p.paragraph_format
    line_spacing_map = {'single': WD_LINE_SPACING.SINGLE, 'one_point_five': WD_LINE_SPACING.ONE_POINT_FIVE, 'double': WD_LINE_SPACING.DOUBLE, 'at_least': WD_LINE_SPACING.AT_LEAST, 'exactly': WD_LINE_SPACING.EXACTLY, 'multiple': WD_LINE_SPACING.MULTIPLE}
    if para_attrs.get("line_spacing_rule") in line_spacing_map:
        pf.line_spacing_rule = line_spacing_map[para_attrs["line_spacing_rule"]]
    if float(para_attrs.get("line_spacing", 0)) > 0:
        pf.line_spacing = Pt(float(para_attrs["line_spacing"]))
    if float(para_attrs.get("space_before", 0)) > 0:
        pf.space_before = Pt(float(para_attrs["space_before"]))
    if float(para_attrs.get("space_after", 0)) > 0:
        pf.space_after = Pt(float(para_attrs["space_after"]))
        
    if float(para_attrs.get("left_indent", 0)) > 0:
        pf.left_indent = Inches(float(para_attrs["left_indent"]) / 1440)
    if float(para_attrs.get("right_indent", 0)) > 0:
        pf.right_indent = Inches(float(para_attrs["right_indent"]) / 1440)
    if float(para_attrs.get("first_line_indent", 0)) > 0:
        pf.first_line_indent = Inches(float(para_attrs["first_line_indent"]) / 1440)

    if para_attrs.get("bottom_border_style") != "none":
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), para_attrs["bottom_border_style"])
        bottom.set(qn('w:sz'), para_attrs.get("bottom_border_size", "4"))
        bottom.set(qn('w:color'), para_attrs.get("bottom_border_color", "auto"))
        pBdr.append(bottom)
        pPr.append(pBdr)

def _add_hyperlink(paragraph, url, runs_data):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    for run_data in runs_data:
        run_element = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if run_data['attrs'].get("bold") == "true": rPr.append(OxmlElement('w:b'))
        r_style = OxmlElement('w:rStyle')
        r_style.set(qn('w:val'), 'Hyperlink')
        rPr.append(r_style)
        run_element.append(rPr)
        t = OxmlElement('w:t')
        t.text = run_data['text']
        run_element.append(t)
        hyperlink.append(run_element)
    paragraph._p.append(hyperlink)

def _create_paragraph_from_xml(p_element, container):
    """Creates and formats a paragraph in the given container (doc or cell)."""
    p = container.add_paragraph()
    _apply_paragraph_properties(p, p_element.attrib)
    
    for child in p_element:
        if child.tag == 'run':
            run = p.add_run()
            text_content = child.find('text').text if child.find('text') is not None else ""
            run.text = text_content
            _apply_run_properties(run, child.attrib)
        elif child.tag == 'hyperlink':
            url = child.get('url')
            runs_data = []
            for run_element in child.findall('run'):
                text_content = run_element.find('text').text if run_element.find('text') is not None else ""
                runs_data.append({'attrs': run_element.attrib, 'text': text_content})
            if url and runs_data: _add_hyperlink(p, url, runs_data)
        elif child.tag == 'image':
            # Image handling needs access to images_data, handled in the main loop
            pass
    return p

def _create_table_from_xml(tbl_element, container, images_data):
    """Creates and formats a table in the given container (doc or cell)."""
    rows = len(tbl_element.findall('row'))
    cols = max(len(row.findall('cell')) for row in tbl_element.findall('row')) if rows > 0 else 0
    if rows == 0 or cols == 0: return
    
    table = container.add_table(rows=rows, cols=cols)
    
    for i, row_element in enumerate(tbl_element.findall('row')):
        for j, cell_element in enumerate(row_element.findall('cell')):
            cell = table.cell(i, j)
            # Clear the default paragraph in the cell
            for p in cell.paragraphs:
                p._p.getparent().remove(p._p)
            # Populate cell with content
            _create_elements_from_xml(cell_element, cell, images_data)

def _create_elements_from_xml(parent_element, container, images_data):
    """Recursively creates elements (paragraphs, tables) in a container."""
    for element in parent_element:
        if element.tag == 'paragraph':
            p = _create_paragraph_from_xml(element, container)
            # Handle images within the paragraph
            for child in element:
                if child.tag == 'image':
                    original_r_id = child.get('r_id')
                    drawing_xml_str = child.get('drawing_xml')
                    if original_r_id in images_data and drawing_xml_str:
                        image_info = images_data[original_r_id]
                        image_stream = BytesIO(image_info['bytes'])
                        try:
                            drawing_element = etree.fromstring(drawing_xml_str)
                            extent = drawing_element.find('.//wp:extent', namespaces=ns)
                            width = Emu(int(extent.get('cx')))
                            height = Emu(int(extent.get('cy')))
                            run = p.add_run()
                            run.add_picture(image_stream, width=width, height=height)
                        except (etree.XMLSyntaxError, AttributeError, ValueError, KeyError):
                            pass # Failsafe
        elif element.tag == 'table_grid':
            _create_table_from_xml(element, container, images_data)

def create_docx_from_xml(optimized_xml_string, images_data):
    """
    Creates a DOCX file from a structured XML string, preserving all captured formatting.
    """
    try:
        xml_match = re.search(r'<resume>.*</resume>', optimized_xml_string, re.DOTALL)
        if not xml_match: raise ValueError("No <resume> tag found in the LLM output.")
        clean_xml = xml_match.group(0)
        root = etree.fromstring(clean_xml)
    except etree.XMLSyntaxError as e:
        raise ValueError(f"Failed to parse XML from LLM output: {e}")

    doc = Document()
    # Clear the default paragraph
    for p in doc.paragraphs:
        p._p.getparent().remove(p._p)
        
    _create_elements_from_xml(root, doc, images_data)
    
    return doc